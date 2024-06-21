from django.shortcuts import render
from google.cloud import storage, vision_v1
from django.http import JsonResponse , HttpResponseNotFound , FileResponse 
from django.conf import settings
from docx import Document
from django.urls import reverse
from PyPDF2 import PdfReader
import io
import json
import re
import os

# this is google api key 
os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = r"D:/files/python_file/Django/OCR_New_Google_Vision_Project/office_GV_key/NewServiceAccountToken.json"


def base(request):

    return render(request, 'base.html')


def Get_file(request):
    # Check if the request method is POST
    if request.method == 'POST':
        # Get the uploaded files and selected language from the request
        files = request.FILES.getlist('files')  # Receive multiple files
        selected_language = request.POST.get('language')
        print("This is get file name: ", files)
        print("Selected language: ", selected_language)
        doc_files = []

        # Upload each file to Google Cloud Storage and process with Google Vision API
        bucket_name = 'ocr_demo1'
        storage_client = storage.Client()
        vision_client = vision_v1.ImageAnnotatorClient()

        for index, file in enumerate(files):
            # Upload the file to Google Cloud Storage
            bucket = storage_client.bucket(bucket_name)
            blob = bucket.blob(file.name)
            blob.upload_from_file(file)

            all_extracted_text = ""

            if file.name.lower().endswith('.pdf'):
                # Process PDF files
                file.seek(0)
                pdf_reader = PdfReader(io.BytesIO(file.read()))
                num_pages = len(pdf_reader.pages)

                timeout_per_page = 60
                total_timeout = num_pages * timeout_per_page

                mime_type = "application/pdf"
                feature = vision_v1.Feature(type=vision_v1.Feature.Type.DOCUMENT_TEXT_DETECTION)

                gcs_source_uri = f'gs://{bucket_name}/{file.name}'
                gcs_source = vision_v1.GcsSource(uri=gcs_source_uri)
                input_config = vision_v1.InputConfig(gcs_source=gcs_source, mime_type=mime_type)

                gcs_destination_uri = f'gs://{bucket_name}/pdf_result/{file.name}'
                gcs_destination = vision_v1.GcsDestination(uri=gcs_destination_uri)
                output_config = vision_v1.OutputConfig(gcs_destination=gcs_destination, batch_size=1)

                async_request = vision_v1.AsyncAnnotateFileRequest(
                    features=[feature], input_config=input_config, output_config=output_config)

                operation = vision_client.async_batch_annotate_files(requests=[async_request])

                print(f"Waiting for operation to complete (timeout: {total_timeout} seconds)...")
                operation.result(timeout=total_timeout)

                match = re.match(r'gs://([^/]+)/(.+)', gcs_destination_uri)
                bucket_name = match.group(1)
                prefix = match.group(2)
                bucket = storage_client.bucket(bucket_name)

                blob_list = list(bucket.list_blobs(prefix=prefix))

                for output in blob_list:
                    json_string = output.download_as_string()

                    response_dict = json.loads(json_string)
                    if "responses" in response_dict and "fullTextAnnotation" in response_dict["responses"][0]:
                        extracted_text = response_dict["responses"][0]["fullTextAnnotation"]["text"]
                        all_extracted_text += extracted_text
                        print(all_extracted_text)
            else:
                # Process image files
                image = vision_v1.Image()
                image.source.image_uri = f'gs://{bucket_name}/{file.name}'

                response = vision_client.document_text_detection(image=image)

                if response.error.message:
                    raise Exception(f'{response.error.message}')

                if 'full_text_annotation' in response:
                    extracted_text = response.full_text_annotation.text
                    all_extracted_text += extracted_text
                    print(all_extracted_text)

            # Save the extracted text to a DOC file
            media_root = settings.MEDIA_ROOT
            if not os.path.exists(media_root):
                os.makedirs(media_root)

            base_name = os.path.splitext(file.name)[0]
            doc_file_path = os.path.join(media_root, f'{base_name}.docx')

            document = Document()
            document.add_heading('OCR Extracted Text', 0)
            document.add_paragraph(all_extracted_text)

            document.save(doc_file_path)
            doc_files.append(doc_file_path)

        # Render the template with a success message and the paths to the DOC files
        return render(request, 'get_file.html', {'success': True, 'doc_files': doc_files})

    return render(request, 'get_file.html')


def download_file(request):
    # Get the list of DOC files in the media folder
    media_folder_path = settings.MEDIA_ROOT
    doc_files = [f for f in os.listdir(media_folder_path) if f.endswith('.doc') or f.endswith('.docx')]
    print("DOC files in media folder:", doc_files)

    download_urls = []
    errors = []

    for doc_file_name in doc_files:
        full_doc_file_path = os.path.join(media_folder_path, doc_file_name)
        print(f"Full Doc File Path: {full_doc_file_path}")

        if os.path.exists(full_doc_file_path):
            # Create a download URL for each DOC file
            download_url = request.build_absolute_uri(reverse('serve_doc_file', args=[doc_file_name]))
            download_urls.append(download_url)
            print(f"Download URL for DOC file: {download_url}")
        else:
            errors.append(f"File does not exist: {full_doc_file_path}")

    if errors:
        # Return error response if any file does not exist
        return JsonResponse({'success': False, 'errors': errors})

    # Return success response with download URLs
    return JsonResponse({'success': True, 'download_urls': download_urls})

def serve_doc_file(request, doc_file_name):
    # Get the full path of the DOC file
    media_folder_path = settings.MEDIA_ROOT
    full_doc_file_path = os.path.join(media_folder_path, doc_file_name)

    if os.path.exists(full_doc_file_path):
        # Create a FileResponse for downloading the DOC file
        response = FileResponse(open(full_doc_file_path, 'rb'), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{doc_file_name}"'

        # Define a function to delete the file after the response is closed
        def delete_file_callback(response):
            try:
                os.remove(full_doc_file_path)
                print(f"File {doc_file_name} has been deleted successfully.")
            except OSError as e:
                print(f"Error: {full_doc_file_path} : {e.strerror}")

        # Attach the callback to the response's close method
        original_close = response.close
        def custom_close(*args, **kwargs):
            original_close(*args, **kwargs)
            delete_file_callback(response)
        response.close = custom_close

        return response
    else:
        # Return 404 response if the file does not exist
        return HttpResponseNotFound('File not found')

# def download_file(request):
#     # Get the list of DOC files in the media folder
#     media_folder_path = settings.MEDIA_ROOT
#     doc_files = [f for f in os.listdir(media_folder_path) if f.endswith('.doc') or f.endswith('.docx')]
#     print("DOC files in media folder:", doc_files)
#     download = request.POST.get('download')
#     print("This is download button request:", download)

#     download_urls = []
#     errors = []

#     for doc_file_name in doc_files:
#         full_doc_file_path = os.path.join(media_folder_path, doc_file_name)
#         print(f"Full Doc File Path: {full_doc_file_path}")

#         if os.path.exists(full_doc_file_path):
#             # Create a download URL for each DOC file
#             download_url = request.build_absolute_uri(reverse('serve_doc_file', args=[doc_file_name]))
#             download_urls.append(download_url)
#             print(f"Download URL for DOC file: {download_url}")
#         else:
#             errors.append(f"File does not exist: {full_doc_file_path}")

#     if errors:
#         # Return error response if any file does not exist
#         return JsonResponse({'success': False, 'errors': errors})

#     # Return success response with download URLs
#     return JsonResponse({'success': True, 'download_urls': download_urls})


# def serve_doc_file(request, doc_file_name):
#     # Get the full path of the DOC file
#     media_folder_path = settings.MEDIA_ROOT
#     full_doc_file_path = os.path.join(media_folder_path, doc_file_name)

#     if os.path.exists(full_doc_file_path):
#         # Create a FileResponse for downloading the DOC file
#         response = FileResponse(open(full_doc_file_path, 'rb'), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#         response['Content-Disposition'] = f'attachment; filename="{doc_file_name}"'

#         # Add a callback to delete the file after the response is completed
#         def delete_file_callback(response):
#             try:
#                 os.remove(full_doc_file_path)
#                 print(f"File {doc_file_name} has been deleted successfully.")
#             except OSError as e:
#                 print(f"Error: {full_doc_file_path} : {e.strerror}")

#         # Add the callback to the response
#         response.close = lambda *args, **kwargs: (
#             FileResponse.close(response, *args, **kwargs),
#             delete_file_callback(response),
#         )

#         return response
#     else:
#         # Return 404 response if the file does not exist
#         return HttpResponseNotFound('File not found')


# def download_file(request):
#     # Get the list of DOC files in the media folder
#     media_folder_path = settings.MEDIA_ROOT
#     doc_files = [f for f in os.listdir(media_folder_path) if f.endswith('.doc') or f.endswith('.docx')]
#     print("DOC files in media folder:", doc_files)

#     download_urls = []
#     errors = []

#     for doc_file_name in doc_files:
#         full_doc_file_path = os.path.join(media_folder_path, doc_file_name)
#         print(f"Full Doc File Path: {full_doc_file_path}")

#         if os.path.exists(full_doc_file_path):
#             # Create a download URL for each DOC file
#             download_url = request.build_absolute_uri(reverse('serve_doc_file', args=[doc_file_name]))
#             download_urls.append(download_url)
#             print(f"Download URL for DOC file: {download_url}")
#         else:
#             errors.append(f"File does not exist: {full_doc_file_path}")

#     if errors:
#         # Return error response if any file does not exist
#         return JsonResponse({'success': False, 'errors': errors})

#     # Return success response with download URLs
#     return JsonResponse({'success': True, 'download_urls': download_urls})


# def serve_doc_file(request, doc_file_name):
#     # Get the full path of the DOC file
#     media_folder_path = settings.MEDIA_ROOT
#     full_doc_file_path = os.path.join(media_folder_path, doc_file_name)

#     if os.path.exists(full_doc_file_path):
#         # Create a FileResponse for downloading the DOC file
#         response = FileResponse(open(full_doc_file_path, 'rb'), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#         response['Content-Disposition'] = f'attachment; filename="{doc_file_name}"'

#         # Add a callback to delete the file after the response is completed
#         def delete_file_callback(response):
#             try:
#                 os.remove(full_doc_file_path)
#                 print(f"File {doc_file_name} has been deleted successfully.")
#             except OSError as e:
#                 print(f"Error: {full_doc_file_path} : {e.strerror}")

#         # Add the callback to the response
#         response.close = lambda *args, **kwargs: (
#             FileResponse.close(response, *args, **kwargs),
#             delete_file_callback(response),
#         )

#         return response
#     else:
#         # Return 404 response if the file does not exist
#         return HttpResponseNotFound('File not found')